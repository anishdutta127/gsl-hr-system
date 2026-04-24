"""
Dump the full text of each .docx template so we can identify variable
slots. Read-only; produces a text-per-template dump to logs/.
"""

from __future__ import annotations

import sys
from pathlib import Path

import docx  # type: ignore

ORIG = Path("public/hr-templates/originals")
OUT = Path("logs/template_inspection.txt")


def main() -> int:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    with OUT.open("w", encoding="utf-8") as f:
        for p in sorted(ORIG.iterdir()):
            if p.suffix.lower() != ".docx":
                continue
            f.write(f"\n\n==================== {p.name} ====================\n")
            try:
                d = docx.Document(str(p))
                for para in d.paragraphs:
                    text = para.text.strip()
                    if text:
                        f.write(text + "\n")
                for table in d.tables:
                    f.write("\n[table]\n")
                    for row in table.rows:
                        cells = [c.text.strip() for c in row.cells]
                        f.write(" | ".join(cells) + "\n")
            except Exception as err:  # noqa: BLE001
                f.write(f"ERROR: {err}\n")
    print(f"Wrote {OUT}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
