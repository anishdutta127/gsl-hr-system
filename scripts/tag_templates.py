"""
Tag the 6 editable .docx templates for docxtemplater consumption.

Reads originals from public/hr-templates/originals/, does targeted
paragraph-level text replacement to swap sample values for {token}
placeholders, writes tagged copies at public/hr-templates/<ID>-v1.docx.

Paragraph-level replacement (vs. run-level) handles the common case
where Word splits a single word across multiple runs. All runs in a
paragraph are joined, substitutions applied, then the full text is
re-written to the first run and other runs cleared. This preserves
the paragraph's style while rewriting its text.

The Experience Letter PDF is a signed sample, not a template, so it
stays in originals/ and is not mirrored.

The Appointment Letter salary structure is a table with sample values
the operator will edit per-offer; treated as text tokens here, not as
computed values. Moving to computed breakdown (Basic = 50% gross, etc.)
is an opt-in future step.
"""

from __future__ import annotations

import shutil
import sys
from pathlib import Path

import docx  # type: ignore
from docx.table import Table  # type: ignore


ORIG = Path("public/hr-templates/originals")
OUT = Path("public/hr-templates")


def replace_in_paragraph(paragraph, pairs: list[tuple[str, str]]) -> None:
    """Join runs, replace, write back to first run, clear others."""
    full = "".join(r.text for r in paragraph.runs)
    if not full:
        return
    replaced = full
    changed = False
    for needle, token in pairs:
        if needle and needle in replaced:
            replaced = replaced.replace(needle, token)
            changed = True
    if not changed:
        return
    if paragraph.runs:
        paragraph.runs[0].text = replaced
        for r in paragraph.runs[1:]:
            r.text = ""


def tag_document(src: Path, dst: Path, pairs: list[tuple[str, str]]) -> int:
    shutil.copy(str(src), str(dst))
    d = docx.Document(str(dst))
    count = 0
    for p in d.paragraphs:
        before = "".join(r.text for r in p.runs)
        replace_in_paragraph(p, pairs)
        after = "".join(r.text for r in p.runs)
        if before != after:
            count += 1
    # Tables too
    for tbl in d.tables:
        if isinstance(tbl, Table):
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        before = "".join(r.text for r in p.runs)
                        replace_in_paragraph(p, pairs)
                        after = "".join(r.text for r in p.runs)
                        if before != after:
                            count += 1
    d.save(str(dst))
    return count


# ------------------------------------------------------------------
# Per-template replacements. Order matters: replace longer strings
# first so a shorter match doesn't eat a piece of a longer one.
# ------------------------------------------------------------------

APPOINTMENT_SALES = ("MTPL Appointment Letter- Sales Team.docx", "APPOINTMENT-SALES-v1.docx", [
    ("23rd April 2026", "{todayLong}"),
    ("Snehasish Mistri", "{name}"),
    ("sketh19@gmail.com", "{email}"),
    ("+91-7003922678", "{phone}"),
    ("Dear Mr. Snehasish", "Dear {title} {firstName}"),
    ("Sr Innovation Mentor", "{designation}"),
    ("STEM & Training", "{department}"),
    ("Kolkata", "{location}"),
    ("Rs. 4,80,000/Annum (Rs. Four Lacs Eighty Thousand Only)", "Rs. {ctcAnnual} (Rs. {ctcWords} Only)"),
    ("24th April 2026", "{joiningDateLong}"),
    ("20,000", "{basicMonthly}"),
    ("240,000", "{basicAnnual}"),
    ("10,000", "{hraMonthly}"),
    ("120,000", "{hraAnnual}"),
    ("40,000", "{grossMonthly}"),
    ("480,000", "{grossAnnual}"),
    ("1,800", "{pfMonthly}"),
    ("21,600", "{pfAnnual}"),
    ("208", "{ptMonthly}"),
    ("2,500", "{ptAnnual}"),
    ("2,008", "{totalDeductionsMonthly}"),
    ("24,100", "{totalDeductionsAnnual}"),
    ("37,992", "{netMonthly}"),
    ("455,900", "{netAnnual}"),
    ("Amit Zaveri", "{signatoryName}"),
    ("Chief Executive Officer", "{signatoryTitle}"),
])

EMPLOYMENT_VERIFICATION = ("Employment Verification Letter.docx", "EMPLOYMENT-VERIFICATION-v1.docx", [
    ("13th November 2025", "{todayLong}"),
    ("Mr. Anish Dutta", "{title} {name}"),
    ("17th June 2025", "{joiningDateLong}"),
    ("Manager – Product", "{designation}"),
    ("Product Department", "{department} Department"),
    ("MTPL/189", "{employeeCode}"),
    ("his/her", "{pronounPossessive}"),
    ("504, C Wing, Chandiwala Pearl Regency, behind reliance digital,", "{addressLine1}"),
    ("Fish Market Area, Navneeth Colony, Andheri West,", "{addressLine2}"),
    ("Mumbai, Maharashtra 400053, India", "{addressLine3}"),
    ("12/07/1997", "{dobShort}"),
    ("Aniruddha Dutta", "{fathersName}"),
    ("932984019360", "{aadhaar}"),
    ("CGQPD8231A", "{pan}"),
    ("Amit Zaveri", "{signatoryName}"),
    ("Chief Executive Officer", "{signatoryTitle}"),
])

INTERNSHIP_COMPLETION = ("MTPL - Internship Completion Letter.docx", "INTERNSHIP-COMPLETION-v1.docx", [
    ("14th June 2025", "{todayLong}"),
    ("Ms. Sheetal Chopde", "{title} {name}"),
    ("sheetalchopde2023@det.sndt.ac.in", "{email}"),
    ("15th April 2025", "{startDateLong}"),
    # Note: the endDate string ("14th June 2025") collides with todayLong above;
    # since todayLong already swapped for the first occurrence, the body's
    # "14th June 2025" (second occurrence) will not be re-matched because the
    # first pass already changed it. Clarify by making end-date a distinct phrase:
    ("to {todayLong}", "to {endDateLong}"),
    ("120 hours", "{hoursCompleted} hours"),
    ("sincere, hardworking, and efficient. She", "sincere, hardworking, and efficient. {pronounSubject}"),
    ("Amit Zaveri", "{signatoryName}"),
    ("Chief Executive Officer", "{signatoryTitle}"),
])

INTERNSHIP_OFFER = ("MTPL Internship Letter.docx", "INTERNSHIP-OFFER-v1.docx", [
    ("07th November 2025", "{todayLong}"),
    ("Mr. Krishna Kanta Misra", "{title} {name}"),
    ("krishnakantamisra8@gmail.com", "{email}"),
    ("Dear Mr. Krishna", "Dear {title} {firstName}"),
    # Swap the second "07th November 2025" that's already been replaced to
    # "{todayLong}" — no, it already became the same token. Use the sentence
    # structure instead:
    ("with effect from {todayLong} to", "with effect from {startDateLong} to"),
    ("06th February 2026", "{endDateLong}"),
    ("Mr. Avishek Dasgupta", "{reportingManager}"),
    ("STEM & Training Department", "{department} Department"),
    ("Rs. 6,000/- (Rupees Five Thousand only)", "Rs. {stipendAmount}/- ({stipendWords} only)"),
    ("Amit Zaveri", "{signatoryName}"),
    ("Chief Executive Officer", "{signatoryTitle}"),
])

RELIEVING = ("MTPL Relieving Letter.docx", "RELIEVING-v1.docx", [
    ("13th April 2026", "{todayLong}"),
    ("Shiva Eppa", "{name}"),
    ("Sr Executive- STEM & Training", "{designation}"),
    ("MTPL139", "{employeeCode}"),
    ("9th September 2024", "{joiningDateLong}"),
    ("14th February 2026", "{lastWorkingDayLong}"),
    ("Amit Zaveri", "{signatoryName}"),
    ("Chief Executive Officer", "{signatoryTitle}"),
])

NO_DUES = ("No Dues Letter Format.docx", "NO-DUES-v1.docx", [
    ("Shweta Mahadik", "{name}"),
    ("MTPL136", "{employeeCode}"),
    ("Rs.104,501", "Rs.{duesAmount}"),
    ("Indian Rupees One lakh Four Thousand Five Hundred and One only", "Indian Rupees {duesAmountWords} only"),
])


def main() -> int:
    specs = [
        EMPLOYMENT_VERIFICATION,
        INTERNSHIP_COMPLETION,
        INTERNSHIP_OFFER,
        APPOINTMENT_SALES,
        RELIEVING,
        NO_DUES,
    ]
    for src_name, dst_name, pairs in specs:
        src = ORIG / src_name
        dst = OUT / dst_name
        if not src.exists():
            print(f"MISSING: {src}", file=sys.stderr)
            continue
        count = tag_document(src, dst, pairs)
        print(f"  {dst.name}: {count} substitutions")
    print("Done.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
